import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from config import (
    DETECTION_MODE,
    DOCUMENT_SIMILARITY_THRESHOLD,
    MAX_SENTENCE_MATCHES_TO_SHOW,
    SENTENCE_EXACT_MATCH_THRESHOLD,
    SENTENCE_MIN_COVERAGE,
    SENTENCE_MIN_EXACT_MATCHES,
    SENTENCE_MIN_LENGTH,
    SENTENCE_MIN_TOTAL_MATCHES,
    SENTENCE_SIMILARITY_THRESHOLD,
    SHOW_DETAILED_SENTENCE_MATCHES,
    SMART_MODE_MIN_SIMILARITY,
)
from constants import COMMON_PHRASES, DEBUG_MODE, END_MARKERS, STOP_WORDS_ES
from sentence_detector import (
    find_matching_sentences,
    format_sentence_match_report,
    is_plagiarism_detected,
)

# --- CONFIGURACIÓN ---
FILES_DIR = Path("./files")
EXTENSIONS = ["*.pdf", "*.docx", "*.md"]


def clean_text(text: str) -> str:
    """Borra consignas repetidas y recorta bibliografía."""

    # Normaliza (elimina espacios múltiples)
    text = " ".join(text.split())

    # 1. Elimina frases repetidas
    for phrase in COMMON_PHRASES:
        # Elimina sin importar el caso
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        text = pattern.sub("", text)

    # 2. Recorta bibliografía
    text_lower = text.lower()
    cut_position = -1

    for marker in END_MARKERS:
        pos = text_lower.rfind(marker)  # Busca la última ocurrencia
        # Si se encuentra y está en el último 20% del documento (para evitar corte falso al inicio)
        if pos != -1 and pos > len(text) * 0.5:
            if cut_position == -1 or pos < cut_position:
                cut_position = pos

    if cut_position != -1:
        text = text[:cut_position]

    return text


def extract_text(file_path: Path) -> Optional[str]:
    """Extrae el texto del archivo."""
    text = ""
    ext = file_path.suffix.lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(str(file_path))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif ext == ".docx":
            doc = Document(str(file_path))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == ".md":
            with file_path.open("r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        print(f"[ERROR] {file_path.name}: {e}")
        return None

    # Aplicamos la limpieza aquí
    return clean_text(text)


def extract_raw_text(file_path: Path) -> Optional[str]:
    """Extrae el texto sin limpiar (para sentence-level detection)."""
    text = ""
    ext = file_path.suffix.lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(str(file_path))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif ext == ".docx":
            doc = Document(str(file_path))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == ".md":
            with file_path.open("r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        print(f"[ERROR] {file_path.name}: {e}")
        return None

    return text


def get_files(directory: Path, extensions: List[str]) -> List[Path]:
    """Busca archivos con las extensiones dadas en el directorio."""
    files = []
    for ext in extensions:
        files.extend(directory.glob(ext))
    return files


def save_debug_text(filename: str, content: str) -> None:
    """Guarda el texto procesado en la carpeta debug."""
    debug_dir = Path("debug")
    if not debug_dir.exists():
        debug_dir.mkdir()

    with (debug_dir / f"{filename}.txt").open("w", encoding="utf-8") as f:
        f.write(content)


def save_results(output_lines: List[str], filename: str = "results.txt") -> None:
    """Guarda los resultados en un archivo de texto dentro de la carpeta output."""
    output_dir = Path("output")
    if not output_dir.exists():
        output_dir.mkdir()

    output_path = output_dir / filename
    with output_path.open("w", encoding="utf-8") as f:
        f.write("\n".join(output_lines))
    print(f"\nℹ️  Los resultados se guardaron en '{output_path}'")


def document_level_similarity(
    texts: List[str], filenames: List[str]
) -> Tuple[List, List]:
    """
    Realiza análisis de similitud TF-IDF a nivel de documento.
    Retorna: (similarity_matrix, suspicious_pairs)
    """
    print("--- Calculando similitudes (TF-IDF) ---")
    vectorizer = TfidfVectorizer(
        stop_words=STOP_WORDS_ES,
        ngram_range=(1, 3),
        min_df=1,
    ).fit_transform(texts)

    vectors = vectorizer.toarray()
    similarity_matrix = cosine_similarity(vectors)

    suspicious_pairs: List[Tuple[str, str, float, str]] = []
    for i in range(len(vectors)):
        for j in range(i + 1, len(vectors)):
            score = similarity_matrix[i][j]
            if score > DOCUMENT_SIMILARITY_THRESHOLD:
                suspicious_pairs.append(
                    (filenames[i], filenames[j], score, "document-level")
                )

    return similarity_matrix, suspicious_pairs


def sentence_level_analysis(
    raw_texts: Dict[str, str], filenames: List[str]
) -> List[Tuple[str, str, float, str, Dict]]:
    """
    Realiza análisis de plagio a nivel de oraciones.
    Retorna: lista de (file1, file2, sentence_score, detection_type, match_stats)
    """
    print("--- Análisis a nivel de oraciones ---")
    suspicious_pairs = []

    total_comparisons = len(filenames) * (len(filenames) - 1) // 2
    comparisons_done = 0

    for i in range(len(filenames)):
        for j in range(i + 1, len(filenames)):
            comparisons_done += 1
            if comparisons_done % 50 == 0:
                print(
                    f"   Progreso: {comparisons_done}/{total_comparisons} comparaciones..."
                )

            file1 = filenames[i]
            file2 = filenames[j]

            match_stats = find_matching_sentences(
                raw_texts[file1],
                raw_texts[file2],
                min_length=SENTENCE_MIN_LENGTH,
                threshold=SENTENCE_SIMILARITY_THRESHOLD,
                exact_threshold=SENTENCE_EXACT_MATCH_THRESHOLD,
            )

            is_plagiarism, reason = is_plagiarism_detected(
                match_stats,
                min_exact_matches=SENTENCE_MIN_EXACT_MATCHES,
                min_total_matches=SENTENCE_MIN_TOTAL_MATCHES,
                min_coverage=SENTENCE_MIN_COVERAGE,
            )

            if is_plagiarism:
                # Usa coverage como pseudo-score para ordenar
                score = match_stats["coverage"]
                suspicious_pairs.append(
                    (file1, file2, score, "sentence-level", match_stats)
                )

    return suspicious_pairs


def smart_mode_analysis(
    texts: List[str],
    raw_texts: Dict[str, str],
    filenames: List[str],
    similarity_matrix: List,
) -> List[Tuple[str, str, float, str, Optional[Dict]]]:
    """
    Smart mode: Realiza análisis por fases.
    - Fase 1: Análisis a nivel de documento (TF-IDF)
    - Fase 2: Análisis a nivel de oraciones (sentence-level)
    Retorna: lista de pares sospechosos
    """
    print("--- Modo inteligente: análisis por fases ---")
    suspicious_pairs = []

    # Fase 1: Document-level (ya hecho, usa similarity_matrix)
    print("   Fase 1: TF-IDF completo")
    candidates_for_phase2 = []

    for i in range(len(filenames)):
        for j in range(i + 1, len(filenames)):
            score = similarity_matrix[i][j]

            if score > DOCUMENT_SIMILARITY_THRESHOLD:
                # Already flagged by document-level
                suspicious_pairs.append(
                    (filenames[i], filenames[j], score, "document-level", None)
                )
            elif SMART_MODE_MIN_SIMILARITY <= score < DOCUMENT_SIMILARITY_THRESHOLD:
                # Gray zone - candidate for sentence-level
                candidates_for_phase2.append((i, j, filenames[i], filenames[j]))

    # Fase 2: Sentence-level en los pares sospechosos
    print(
        f"   Fase 2: Análisis de oraciones en {len(candidates_for_phase2)} pares sospechosos"
    )

    for idx, (i, j, file1, file2) in enumerate(candidates_for_phase2, 1):
        if idx % 10 == 0:
            print(f"      Progreso: {idx}/{len(candidates_for_phase2)}...")

        match_stats = find_matching_sentences(
            raw_texts[file1],
            raw_texts[file2],
            min_length=SENTENCE_MIN_LENGTH,
            threshold=SENTENCE_SIMILARITY_THRESHOLD,
            exact_threshold=SENTENCE_EXACT_MATCH_THRESHOLD,
        )

        is_plagiarism, reason = is_plagiarism_detected(
            match_stats,
            min_exact_matches=SENTENCE_MIN_EXACT_MATCHES,
            min_total_matches=SENTENCE_MIN_TOTAL_MATCHES,
            min_coverage=SENTENCE_MIN_COVERAGE,
        )

        if is_plagiarism:
            pseudo_score = match_stats["coverage"]
            suspicious_pairs.append(
                (file1, file2, pseudo_score, "sentence-level", match_stats)
            )

    return suspicious_pairs


def select_detection_mode() -> str:
    """
    Menu interactivo para seleccionar el modo de detección.
    Retorna el modo seleccionado o el default de config.
    """
    print(f"\n{'=' * 80}")
    print("DETECTOR DE PLAGIO - Selección de Modo")
    print(f"{'=' * 80}\n")

    print("Modos disponibles:\n")
    print("  1. FAST      - Solo TF-IDF (⚡ ~segundos, detecta copias directas)")
    print(
        "  2. THOROUGH  - Solo análisis de oraciones (🔍 ~1-2 min, detecta plagio sofisticado)"
    )
    print("  3. HYBRID    - Ambos análisis (🎯 ~2-3 min, máxima precisión)")
    print(
        "  4. SMART     - Inteligente en 2 fases (🧠 ~20 seg, balanceado) [RECOMENDADO]"
    )

    print(f"\nModo por defecto: {DETECTION_MODE.upper()}")
    print(
        "Presioná Enter para usar el modo por defecto, o escribí el número/nombre del modo:\n"
    )
    print("\nO escribí '0' o 'exit' para salir\n")

    user_input = (
        input("Selección (0/1-4 o fast/thorough/hybrid/smart/exit): ").strip().lower()
    )

    # Si el usuario escribe '0' o 'exit', sale
    if user_input in ["0", "exit"]:
        print("\nSaliendo del script, ¡hasta luego!\n")
        exit()

    # Si el usuario escribe Enter, usa el modo por defecto
    if not user_input:
        return DETECTION_MODE

    # Mapeo de números a nombres de modos
    mode_map = {
        "1": "fast",
        "2": "thorough",
        "3": "hybrid",
        "4": "smart",
        "fast": "fast",
        "thorough": "thorough",
        "hybrid": "hybrid",
        "smart": "smart",
    }

    selected_mode = mode_map.get(user_input)

    if selected_mode:
        print(f"\n✅ Modo seleccionado: {selected_mode.upper()}\n")
        return selected_mode
    else:
        print(
            f"\n⚠️  Opción inválida. Usando modo por defecto: {DETECTION_MODE.upper()}\n"
        )
        return DETECTION_MODE


def analyze_similarities() -> None:
    """
    Analiza las similitudes entre los archivos según el modo configurado.
    """
    # Interactive mode selection
    mode = select_detection_mode()

    print(f"{'=' * 80}")
    print(f"DETECTOR DE PLAGIO - Modo: {mode.upper()}")
    print(f"{'=' * 80}")

    files = get_files(FILES_DIR, EXTENSIONS)

    if not files:
        print("No se encontraron archivos. Tenes que ponerlos en la carpeta 'files'")
        return

    print(
        f"Procesando {len(files)} archivos (Limpieza de frases y bibliografía y cosas así)..."
    )

    if DEBUG_MODE:
        print("[DEBUG] Guardando textos limpiados en 'debug/' folder")

    texts: List[str] = []
    raw_texts: Dict[str, str] = {}
    filenames: List[str] = []

    for file in files:
        content = extract_text(file)
        raw_content = extract_raw_text(file)

        # Filtra los textos que están vacíos después de la limpieza
        if content and len(content.strip()) > 50:
            texts.append(content)
            filename = file.name
            filenames.append(filename)

            if raw_content:
                raw_texts[filename] = raw_content

            if DEBUG_MODE:
                save_debug_text(filename, content)

    if len(texts) < 2:
        print("At least 2 valid files are needed to compare.")
        return

    suspicious_pairs = []

    # Ejecuta según el modo
    if mode == "fast":
        # Modo rápido: Solo TF-IDF
        _, suspicious_pairs = document_level_similarity(texts, filenames)

    elif mode == "thorough":
        # Modo exhaustivo: Solo análisis por oraciones
        suspicious_pairs = sentence_level_analysis(raw_texts, filenames)

    elif mode == "hybrid":
        # Modo híbrido: Ambos análisis
        similarity_matrix, doc_pairs = document_level_similarity(texts, filenames)
        sent_pairs = sentence_level_analysis(raw_texts, filenames)

        # Combina ambos, evitando duplicados
        suspicious_pairs = doc_pairs
        for sent_pair in sent_pairs:
            # Check if not already in doc_pairs
            file1, file2 = sent_pair[0], sent_pair[1]
            is_duplicate = any(
                (p[0] == file1 and p[1] == file2) or (p[0] == file2 and p[1] == file1)
                for p in doc_pairs
            )
            if not is_duplicate:
                suspicious_pairs.append(sent_pair)

    elif mode == "smart":
        # Modo inteligente: Enfoque en dos fases
        similarity_matrix, _ = document_level_similarity(texts, filenames)
        suspicious_pairs = smart_mode_analysis(
            texts, raw_texts, filenames, similarity_matrix
        )

    else:
        print(f"❌ Modo desconocido: {mode}")
        print("   Modos válidos: fast, thorough, hybrid, smart")
        return

    # Ordena por score
    suspicious_pairs.sort(key=lambda x: x[2], reverse=True)

    # Formatea la salida
    output_lines = []
    output_lines.append(f"\nResultados de análisis - Modo: {mode.upper()}")
    output_lines.append("=" * 80)

    if not suspicious_pairs:
        output_lines.append("✅ No se detectaron copias obvias.")
    else:
        output_lines.append(
            f"⚠️  Se encontraron {len(suspicious_pairs)} pares sospechosos:\n"
        )

        for item in suspicious_pairs:
            f1, f2, score, detection_type = item[0], item[1], item[2], item[3]
            match_stats = item[4] if len(item) > 4 else None

            if detection_type == "document-level":
                output_lines.append(f"🔴 {score * 100:.2f}% :: {f1} <--> {f2}")
                output_lines.append(
                    "   📄 Detectado por: TF-IDF (similitud documental)"
                )
            else:  # sentence-level
                output_lines.append(f"🔴 Detectado :: {f1} <--> {f2}")
                output_lines.append("   📝 Detectado por: Análisis de oraciones")

                if match_stats and SHOW_DETAILED_SENTENCE_MATCHES:
                    report = format_sentence_match_report(
                        f1, f2, match_stats, MAX_SENTENCE_MATCHES_TO_SHOW
                    )
                    output_lines.extend(report)

            output_lines.append("")  # Línea vacía entre pares

    # Imprime en consola
    for line in output_lines:
        try:
            print(line)
        except UnicodeEncodeError:
            print(line.encode("ascii", "replace").decode("ascii"))

    # Guarda en archivo
    save_results(
        output_lines,
        "results_" + mode + datetime.now().strftime("_%Y%m%d_%H%M%S") + ".txt",
    )

    # Limpieza del debug si se requiere
    if not DEBUG_MODE and Path("debug").exists():
        shutil.rmtree("debug")


if __name__ == "__main__":
    analyze_similarities()
